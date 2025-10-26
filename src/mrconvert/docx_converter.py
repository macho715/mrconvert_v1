from __future__ import annotations

from pathlib import Path
from typing import Dict, Any, List

try:
    import mammoth  # Preferred for DOCX→HTML/Markdown
except Exception:  # pragma: no cover
    mammoth = None

try:
    import docx  # python-docx fallback
except Exception:  # pragma: no cover
    docx = None

from .utils import now_iso

def _docx_to_markdown_with_mammoth(path: Path) -> str | None:
    if mammoth is None:
        return None
    with open(path, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    return result.value  # type: ignore[attr-defined]

def _docx_plain_with_python_docx(path: Path) -> Dict[str, Any]:
    if docx is None:
        return {"text": "", "tables": []}
    d = docx.Document(str(path))
    paras = [p.text for p in d.paragraphs]
    text = "\n".join([t for t in paras if t.strip()])
    tables: List[Dict[str, Any]] = []
    for ti, tbl in enumerate(d.tables):
        rows = []
        for r in tbl.rows:
            cells = [c.text for c in r.cells]
            rows.append(cells)
        tables.append({"page": None, "index": ti, "rows": rows})
    return {"text": text, "tables": tables}

def extract_docx(path: Path) -> Dict[str, Any]:
    md = _docx_to_markdown_with_mammoth(path)
    plain_tables = _docx_plain_with_python_docx(path)
    return {
        "meta": {
            "source": str(path),
            "type": "docx",
            "pages": None,
            "parsed_at": now_iso(),
            "ocr": {"used": False, "engine": "none", "lang": None},
        },
        "text": plain_tables.get("text", ""),
        "markdown": md,
        "tables": plain_tables.get("tables", []),
    }
